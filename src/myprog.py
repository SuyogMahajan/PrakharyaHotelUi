import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document
import time
import random

# Initialize Word document
doc = Document()

# Headers to mimic browser behavior
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/110.0.5481.104 Safari/537.36'
}

output_file = 'examtopics_questions_updated.docx'

def save_document():
    """Saves the Word document."""
    try:
        doc.save(output_file)
        print(f"Document saved successfully: {output_file}")
    except Exception as e:
        print(f"Failed to save document: {e}")

# Loop through question numbers 1 to 5 (expand to full 1-378 after testing)
for question_num in range(1, 6):
    try:
        # Construct Bing search query URL
        query = f'examtopics professional data engineer question {question_num}'
        search_url = f'https://www.bing.com/search?q={query.replace(" ", "+")}'
        
        print(f"Searching: {search_url}")

        # Request the Bing search results page
        response = requests.get(search_url, headers=headers)
        print(f"status {response.status_code}" )
        if response.status_code != 200:
            print(f"Failed to fetch Bing search results for question {question_num}.")
            continue

        soup = BeautifulSoup(response.text, 'html.parser')

        # Find the first search result containing 'examtopics.com'
        first_result = None
        for a_tag in soup.find_all('a', href=True):
            if 'examtopics.com' in a_tag['href']:
                first_result = a_tag['href']
                break

        if not first_result:
            print(f"No valid examtopics.com link found for question {question_num}.")
            continue

        print(f"Fetching content from: {first_result}")

        # Fetch the content of the first result page
        page_response = requests.get(first_result, headers=headers)
        if page_response.status_code != 200:
            print(f"Failed to fetch page content for question {question_num}.")
            continue

        # Parse the fetched page
        page_soup = BeautifulSoup(page_response.text, 'html.parser')

        # Target specific HTML structure (update selector based on actual HTML structure)
        question_content = page_soup.select_one('.question-content')  # Adjust selector as needed
        page_content = question_content.get_text(separator='\n').strip() if question_content else "Content not found."

        if not page_content:
            print(f"No content found on the page for question {question_num}.")
            continue

        # Add question header and content to Word document
        doc.add_heading(f'Question #: {question_num}', level=1)
        doc.add_paragraph(page_content)

        # Add source link as a comment
        doc.add_paragraph(f"Source: {first_result}")

        print(f"Content for question {question_num} added.")

        # Save periodically to avoid data loss
        if question_num % 10 == 0:
            save_document()

        # Delay to prevent blocking (randomized)
        time.sleep(random.uniform(2, 5))

    except Exception as e:
        print(f"Error processing question {question_num}: {e}")

# Save the final document
save_document()
